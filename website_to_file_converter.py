# bash
# brew install --cask libreoffice
# libreoffice --headless --convert-to doc --outdir /path/to/output /path/to/your/document.docx

import subprocess
import requests
from bs4 import BeautifulSoup
from docx import Document
import pdfkit


# Convert Website to PDF using pdfkit
def convert_to_pdf(url, output_path):
    pdfkit.from_url(url, output_path)


# Convert Website to DOCX using BeautifulSoup and python-docx
def convert_to_docx(url, output_path):
    response = requests.get(url)
    soup = BeautifulSoup(response.text, "html.parser")

    doc = Document()
    title = soup.title.string if soup.title else "Website Content"
    doc.add_heading(title, 0)

    paragraphs = soup.find_all('p')
    for para in paragraphs:
        doc.add_paragraph(para.get_text())

    doc.save(output_path)


# Convert DOCX to DOC using LibreOffice command-line tool
def convert_docx_to_doc(docx_path, output_dir):
    subprocess.run(['libreoffice', '--headless', '--convert-to', 'doc', '--outdir', output_dir, docx_path])


# Main Function
def convert_website_to_files(url, pdf_output, docx_output, doc_output_dir):
    # Step 1: Convert Website to PDF
    convert_to_pdf(url, pdf_output)
    print(f"PDF conversion complete: {pdf_output}")

    # Step 2: Convert Website to DOCX
    convert_to_docx(url, docx_output)
    print(f"DOCX conversion complete: {docx_output}")

    # Step 3: Convert DOCX to DOC
    convert_docx_to_doc(docx_output, doc_output_dir)
    print(f"DOC conversion complete. Saved in: {doc_output_dir}")


# Example usage
url = "https://entrust.wd1.myworkdayjobs.com/en-US/EntrustCareers/job/United-Kingdom---London-Onfido/Intern-Applied-Scientist---UK_R003336"
pdf_output = "output.pdf"
docx_output = "output.docx"
doc_output_dir = "/path/to/output"

convert_website_to_files(url, pdf_output, docx_output, doc_output_dir)
